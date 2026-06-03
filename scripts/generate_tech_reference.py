"""Generate HTML and DOCX from docs/tutorials/MCP_TECHNICAL_REFERENCE.md.

Usage:
    python scripts/generate_tech_reference.py

Produces:
    docs/tutorials/MCP_TECHNICAL_REFERENCE.html
    docs/tutorials/MCP_TECHNICAL_REFERENCE.docx (requires python-docx)
"""

import re
import sys
from datetime import datetime, timezone
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SRC = PROJECT_ROOT / 'docs' / 'tutorials' / 'MCP_TECHNICAL_REFERENCE.md'
OUT_DIR = PROJECT_ROOT / 'docs' / 'tutorials'


# ---------------------------------------------------------------------------
# HTML generation
# ---------------------------------------------------------------------------

def generate_html(md_text: str, output: Path) -> None:
	"""Convert markdown to a styled, self-contained HTML file."""
	import markdown

	body = markdown.markdown(
		md_text,
		extensions=['tables', 'fenced_code', 'codehilite', 'toc'],
		extension_configs={
			'toc': {'permalink': True},
			'codehilite': {'css_class': 'highlight', 'guess_lang': False},
		},
	)

	timestamp = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')

	html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>MCP Technical Reference</title>
<style>
body {{
	font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
	max-width: 960px;
	margin: 40px auto;
	padding: 0 24px;
	color: #1a1a1a;
	line-height: 1.65;
	font-size: 16px;
}}
h1 {{
	color: #1a5276;
	border-bottom: 3px solid #1a5276;
	padding-bottom: 12px;
	margin-top: 48px;
}}
h2 {{
	color: #2c3e50;
	border-bottom: 1px solid #ddd;
	padding-bottom: 6px;
	margin-top: 40px;
}}
h3 {{ color: #34495e; margin-top: 28px; }}
h4 {{ color: #555; margin-top: 20px; }}
a {{ color: #2874a6; }}
a.headerlink {{ font-size: 0.8em; color: #bbb; padding-left: 6px; text-decoration: none; }}
a.headerlink:hover {{ color: #2874a6; }}
table {{
	border-collapse: collapse;
	width: 100%;
	margin: 16px 0 24px 0;
	font-size: 0.95em;
}}
th, td {{
	border: 1px solid #ddd;
	padding: 10px 12px;
	text-align: left;
}}
th {{ background-color: #f2f8fc; font-weight: 600; }}
tr:nth-child(even) {{ background-color: #fafafa; }}
code {{
	background: #f0f0f0;
	padding: 2px 6px;
	border-radius: 3px;
	font-size: 0.9em;
	font-family: 'Consolas', 'Menlo', 'Monaco', monospace;
}}
pre {{
	background: #f5f5f5;
	border: 1px solid #ddd;
	border-radius: 4px;
	padding: 14px 18px;
	overflow-x: auto;
	font-size: 0.88em;
	line-height: 1.5;
}}
pre code {{
	background: none;
	padding: 0;
}}
blockquote {{
	border-left: 4px solid #2874a6;
	margin: 16px 0;
	padding: 8px 16px;
	color: #555;
	background: #f9fbfd;
}}
hr {{
	border: none;
	border-top: 2px solid #e0e0e0;
	margin: 36px 0;
}}
.toc {{
	background: #f8f9fa;
	border: 1px solid #e0e0e0;
	border-radius: 6px;
	padding: 16px 24px;
	margin: 24px 0;
}}
.toc ul {{ padding-left: 20px; }}
.toc li {{ margin: 4px 0; }}
.timestamp {{
	color: #999;
	font-size: 0.85em;
	margin-top: -8px;
}}
</style>
</head>
<body>
<p class="timestamp">Generated: {timestamp}</p>
{body}
</body>
</html>"""

	output.write_text(html, encoding='utf-8')
	print(f"  HTML: {output}")


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def generate_docx(md_text: str, output: Path) -> None:
	"""Convert markdown to DOCX with basic formatting."""
	try:
		from docx import Document
		from docx.shared import Pt, Inches, RGBColor
		from docx.enum.text import WD_ALIGN_PARAGRAPH
	except ImportError:
		print("  python-docx not installed. Skipping DOCX.")
		print("  Install with: pip install python-docx")
		return

	doc = Document()
	timestamp = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')

	# Style adjustments
	style = doc.styles['Normal']
	style.font.name = 'Calibri'
	style.font.size = Pt(11)
	style.paragraph_format.space_after = Pt(6)

	lines = md_text.split('\n')
	i = 0
	in_code_block = False
	code_buffer = []
	in_table = False
	table_rows = []

	def flush_table():
		nonlocal table_rows, in_table
		if not table_rows:
			return
		# Parse header and data rows
		headers = [c.strip() for c in table_rows[0].strip('|').split('|')]
		data = []
		for row in table_rows[2:]:  # skip separator row
			cells = [c.strip() for c in row.strip('|').split('|')]
			data.append(cells)

		ncols = len(headers)
		tbl = doc.add_table(rows=1 + len(data), cols=ncols)
		tbl.style = 'Light List Accent 1'
		for ci, h in enumerate(headers):
			tbl.rows[0].cells[ci].text = h
		for ri, row_data in enumerate(data):
			for ci in range(min(ncols, len(row_data))):
				tbl.rows[ri + 1].cells[ci].text = row_data[ci]

		table_rows = []
		in_table = False

	def add_formatted_paragraph(text, style_name=None):
		"""Add a paragraph with inline code/bold/italic formatting."""
		p = doc.add_paragraph(style=style_name)
		# Split on inline code, bold, and italic patterns
		parts = re.split(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))', text)
		for part in parts:
			if part.startswith('`') and part.endswith('`'):
				run = p.add_run(part[1:-1])
				run.font.name = 'Consolas'
				run.font.size = Pt(9.5)
				run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
			elif part.startswith('**') and part.endswith('**'):
				run = p.add_run(part[2:-2])
				run.bold = True
			elif part.startswith('*') and part.endswith('*'):
				run = p.add_run(part[1:-1])
				run.italic = True
			elif part.startswith('['):
				# Link: extract text
				m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
				if m:
					run = p.add_run(m.group(1))
					run.font.color.rgb = RGBColor(0x28, 0x74, 0xa6)
					run.underline = True
				else:
					p.add_run(part)
			else:
				p.add_run(part)
		return p

	while i < len(lines):
		line = lines[i]

		# Code blocks
		if line.strip().startswith('```'):
			if in_code_block:
				# End code block
				code_text = '\n'.join(code_buffer)
				p = doc.add_paragraph()
				run = p.add_run(code_text)
				run.font.name = 'Consolas'
				run.font.size = Pt(9)
				p.paragraph_format.space_before = Pt(6)
				p.paragraph_format.space_after = Pt(6)
				code_buffer = []
				in_code_block = False
			else:
				flush_table()
				in_code_block = True
			i += 1
			continue

		if in_code_block:
			code_buffer.append(line)
			i += 1
			continue

		# Tables
		if '|' in line and line.strip().startswith('|'):
			if not in_table:
				in_table = True
				table_rows = []
			table_rows.append(line)
			i += 1
			continue
		elif in_table:
			flush_table()

		stripped = line.strip()

		# Blank lines
		if not stripped:
			i += 1
			continue

		# Horizontal rules
		if stripped in ('---', '***', '___'):
			doc.add_paragraph().paragraph_format.space_after = Pt(12)
			i += 1
			continue

		# Headings
		if stripped.startswith('#'):
			m = re.match(r'^(#{1,4})\s+(.+)', stripped)
			if m:
				level = len(m.group(1))
				text = m.group(2)
				# Remove anchor links
				text = re.sub(r'\[([^\]]+)\]\(#[^)]+\)', r'\1', text)
				doc.add_heading(text, level=min(level, 4))
				i += 1
				continue

		# List items
		if re.match(r'^[-*]\s', stripped):
			text = re.sub(r'^[-*]\s+', '', stripped)
			add_formatted_paragraph(text, 'List Bullet')
			i += 1
			continue

		if re.match(r'^\d+\.\s', stripped):
			text = re.sub(r'^\d+\.\s+', '', stripped)
			add_formatted_paragraph(text, 'List Number')
			i += 1
			continue

		# Regular paragraph
		add_formatted_paragraph(stripped)
		i += 1

	# Flush any remaining table
	flush_table()

	# Add generation timestamp at top
	doc.paragraphs[0].insert_paragraph_before(
		f'Generated: {timestamp}'
	).style = doc.styles['Normal']

	doc.save(str(output))
	print(f"  DOCX: {output}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
	if not SRC.exists():
		print(f"ERROR: Source not found: {SRC}")
		sys.exit(1)

	print("MCP Technical Reference — Document Generator")
	print("=" * 48)

	md_text = SRC.read_text(encoding='utf-8')
	print(f"  Source: {SRC} ({len(md_text):,} chars)")

	html_out = OUT_DIR / 'MCP_TECHNICAL_REFERENCE.html'
	docx_out = OUT_DIR / 'MCP_TECHNICAL_REFERENCE.docx'

	generate_html(md_text, html_out)
	generate_docx(md_text, docx_out)

	print("\nDone.")


if __name__ == '__main__':
	main()
