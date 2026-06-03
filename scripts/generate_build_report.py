"""Generate process documentation report (HTML + DOCX) for the LTEM Database MCP Server.

Usage:
    python scripts/generate_build_report.py

Produces:
    docs/build_report.html
    docs/build_report.docx (requires python-docx)
"""

import json
import os
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

# Ensure project root is on sys.path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


def get_version() -> str:
	"""Read version from pyproject.toml."""
	toml_path = PROJECT_ROOT / "pyproject.toml"
	for line in toml_path.read_text(encoding="utf-8").splitlines():
		if line.strip().startswith("version"):
			# version = "1.0.0"
			return line.split("=", 1)[1].strip().strip('"').strip("'")
	return "unknown"


def get_git_hash() -> str:
	"""Get short git commit hash, or 'unknown' if not in a git repo."""
	try:
		result = subprocess.run(
			["git", "rev-parse", "--short", "HEAD"],
			capture_output=True, text=True, cwd=str(PROJECT_ROOT),
			timeout=5,
		)
		if result.returncode == 0:
			return result.stdout.strip()
	except Exception:
		pass
	return "unknown"


def get_schema_info() -> dict | None:
	"""Try to connect and get live schema info. Returns None on failure."""
	try:
		from mcp_server.schema import build_schema_snapshot
		return build_schema_snapshot()
	except Exception as e:
		print(f"  Warning: Could not connect to database: {e}")
		return None


def get_tool_catalog() -> list[dict]:
	"""Return catalog of all MCP tools with descriptions."""
	return [
		# Core
		{"module": "Core", "name": "health_check", "description": "Verify database connectivity and return server info"},
		{"module": "Core", "name": "list_tables", "description": "List all tables in the ecological_monitoring database"},
		{"module": "Core", "name": "describe_table", "description": "Column info for a specific table"},
		{"module": "Core", "name": "schema_snapshot", "description": "Full schema with tables, columns, types, row counts"},
		# Data Access
		{"module": "Data Access", "name": "get_regions", "description": "List all surveyed regions"},
		{"module": "Data Access", "name": "get_reefs", "description": "List reefs, optionally filtered by region"},
		{"module": "Data Access", "name": "get_species_list", "description": "Species observed with optional filters"},
		{"module": "Data Access", "name": "get_observations", "description": "Raw data query with filters and limit"},
		{"module": "Data Access", "name": "survey_effort_summary", "description": "Aggregated survey counts by Year/Region/MPA"},
		# Fish Community
		{"module": "Fish Community", "name": "calculate_diversity", "description": "Shannon, Simpson, Pielou diversity indices per survey unit"},
		{"module": "Fish Community", "name": "species_composition", "description": "Top N species by relative abundance per group"},
		{"module": "Fish Community", "name": "trophic_structure", "description": "Biomass proportions by trophic group"},
		{"module": "Fish Community", "name": "size_structure", "description": "Abundance by size class (cm)"},
		{"module": "Fish Community", "name": "community_comparison", "description": "Bray-Curtis dissimilarity matrix between groups"},
		# Biomass
		{"module": "Biomass", "name": "biomass_by_region", "description": "Mean biomass per region + Kruskal-Wallis test"},
		{"module": "Biomass", "name": "biomass_by_depth", "description": "Shallow vs Deep biomass comparison"},
		{"module": "Biomass", "name": "trophic_biomass", "description": "Biomass by trophic group"},
		{"module": "Biomass", "name": "environmental_correlations", "description": "Spearman: biomass vs SST/Chl-a"},
		{"module": "Biomass", "name": "sst_biomass_relationship", "description": "Linear + quadratic SST regression"},
		{"module": "Biomass", "name": "chl_productivity_relationship", "description": "Log-log Chl-a vs productivity"},
		{"module": "Biomass", "name": "latitudinal_gradient", "description": "Biomass trends by latitude"},
		# MPA Effectiveness
		{"module": "MPA Effectiveness", "name": "compare_protection_levels", "description": "Kruskal-Wallis + pairwise Mann-Whitney across protection levels"},
		{"module": "MPA Effectiveness", "name": "cabo_pulmo_recovery", "description": "Recovery trajectory + recovery factor vs baseline"},
		{"module": "MPA Effectiveness", "name": "compare_all_metrics", "description": "Multi-metric comparison with CP advantage %"},
		{"module": "MPA Effectiveness", "name": "trophic_comparison", "description": "Trophic proportions by protection level"},
		{"module": "MPA Effectiveness", "name": "size_comparison", "description": "Size class proportions, large fish >40cm %"},
		{"module": "MPA Effectiveness", "name": "baci_analysis", "description": "Before-After-Control-Impact analysis"},
		{"module": "MPA Effectiveness", "name": "spillover_analysis", "description": "Biomass by distance from Cabo Pulmo"},
		# Temporal Trends
		{"module": "Temporal Trends", "name": "annual_time_series", "description": "Annual aggregated time series"},
		{"module": "Temporal Trends", "name": "trend_analysis", "description": "Linear regression + Mann-Kendall + Sen's slope"},
		{"module": "Temporal Trends", "name": "regional_trends", "description": "Trend comparison across all regions"},
		{"module": "Temporal Trends", "name": "change_point_detection", "description": "Pettitt or CUSUM change point analysis"},
		{"module": "Temporal Trends", "name": "seasonal_patterns", "description": "Monthly aggregation"},
		{"module": "Temporal Trends", "name": "moving_window", "description": "Rolling average + smoothing"},
	]


def get_resource_catalog() -> list[dict]:
	"""Return catalog of MCP resources."""
	return [
		{"uri": "ltem://schema", "description": "Full DB schema (dynamic)"},
		{"uri": "ltem://regions", "description": "LTEM survey regions with metadata"},
		{"uri": "ltem://sampling-protocol", "description": "Sampling design: belt transects 50m x 5m, 4 replicates, 2 depths"},
		{"uri": "ltem://protection-categories", "description": "Protection status definitions"},
		{"uri": "ltem://trophic-groups", "description": "Trophic group cutoffs"},
	]


def generate_html(schema: dict | None, version: str = "unknown", git_hash: str = "unknown") -> str:
	"""Generate HTML report content."""
	timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
	tools = get_tool_catalog()
	resources = get_resource_catalog()

	# Group tools by module
	modules: dict[str, list] = {}
	for t in tools:
		modules.setdefault(t["module"], []).append(t)

	tools_html = ""
	for mod, mod_tools in modules.items():
		tools_html += f"<h3>{mod}</h3>\n<table><tr><th>Tool</th><th>Description</th></tr>\n"
		for t in mod_tools:
			tools_html += f"<tr><td><code>{t['name']}</code></td><td>{t['description']}</td></tr>\n"
		tools_html += "</table>\n"

	resources_html = "<table><tr><th>URI</th><th>Description</th></tr>\n"
	for r in resources:
		resources_html += f"<tr><td><code>{r['uri']}</code></td><td>{r['description']}</td></tr>\n"
	resources_html += "</table>\n"

	schema_html = "<p>Could not connect to database.</p>"
	if schema and "tables" in schema:
		schema_html = ""
		for tbl, info in schema["tables"].items():
			if "error" in info:
				schema_html += f"<h3>{tbl}</h3><p>Error: {info['error']}</p>\n"
				continue
			schema_html += f"<h3>{tbl} ({info.get('row_count', '?')} rows)</h3>\n"
			schema_html += "<table><tr><th>Column</th><th>Type</th><th>Nullable</th></tr>\n"
			for col in info.get("columns", []):
				schema_html += (
					f"<tr><td>{col['name']}</td><td>{col['type']}</td>"
					f"<td>{'Yes' if col.get('nullable') else 'No'}</td></tr>\n"
				)
			schema_html += "</table>\n"

	html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>LTEM Database MCP Server — Build Report</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 900px; margin: 40px auto; padding: 0 20px; color: #333; }}
h1 {{ color: #1a5276; border-bottom: 2px solid #1a5276; padding-bottom: 10px; }}
h2 {{ color: #2c3e50; margin-top: 30px; }}
h3 {{ color: #34495e; }}
table {{ border-collapse: collapse; width: 100%; margin: 10px 0 20px 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background-color: #f2f8fc; }}
code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
.warning {{ color: #e67e22; }}
.timestamp {{ color: #999; font-size: 0.9em; }}
</style>
</head>
<body>

<h1>LTEM Database MCP Server — Build Report</h1>
<p class="timestamp">Version: {version} | Git: {git_hash} | Generated: {timestamp}</p>

<h2>1. Project Goal and Scope</h2>
<p>A FastMCP server connecting to the AWS RDS MySQL <code>ecological_monitoring</code> database,
exposing the full LTEM (Long-Term Ecological Monitoring) analysis suite as MCP tools for
Claude and chatMPA Studio.</p>
<p><strong>In scope:</strong> Database queries, ecological analysis tools, MCP resources, structured JSON responses.</p>
<p><strong>Out of scope:</strong> Visualizations, PDF parsing, embeddings, data cleaning, deployment automation.</p>

<h2>2. Security Model</h2>
<ul>
<li>Read-only database user (<code>mcp_ltem_ro</code>, SELECT only)</li>
<li>Credentials via <code>.env</code> file (never hardcoded, gitignored)</li>
<li>SQL validation: only SELECT/SHOW/DESCRIBE allowed</li>
<li>Denied keywords: INSERT, UPDATE, DELETE, ALTER, DROP, CREATE, TRUNCATE, etc.</li>
<li>Table whitelist (3 core tables)</li>
<li>Auto LIMIT cap (5000 rows)</li>
<li>Query timeout (20 seconds)</li>
<li>Parameterized queries for all user-facing tools</li>
</ul>

<h2>3. Repository Structure</h2>
<pre>
sandbox/ltem-db-mcp/
+-- .env.example          # Credential template
+-- .env                  # Actual credentials (gitignored)
+-- .gitignore
+-- pyproject.toml        # Dependencies
+-- mcp_server/
|   +-- server.py         # FastMCP entry point
|   +-- db.py             # DB connection layer
|   +-- security.py       # SQL validation
|   +-- schema.py         # Schema discovery
+-- tools/
|   +-- data_access.py    # Raw data queries
|   +-- fish_community.py # Diversity, composition
|   +-- biomass.py        # Biomass analysis
|   +-- mpa_effectiveness.py  # MPA comparisons
|   +-- temporal_trends.py    # Time series
+-- scripts/
|   +-- dbconnect-source.R    # (DEPRECATED)
|   +-- generate_build_report.py
+-- docs/
</pre>

<h2>4. Implementation Summary</h2>
<ol>
<li>Project scaffolding (pyproject.toml, .env, .gitignore)</li>
<li>Security layer (SQL validation, table whitelist, LIMIT enforcement)</li>
<li>Database connection layer (PyMySQL, parameterized queries)</li>
<li>Schema discovery with in-memory caching</li>
<li>FastMCP server with 5 resources and 30+ tools</li>
<li>Data access tools (regions, reefs, species, observations)</li>
<li>Fish community tools (diversity, composition, trophic, size)</li>
<li>Biomass tools (regional, depth, environmental correlations)</li>
<li>MPA effectiveness tools (protection comparison, Cabo Pulmo recovery, BACI)</li>
<li>Temporal trend tools (Mann-Kendall, change points, rolling averages)</li>
</ol>

<h2>5. Schema Discovery</h2>
{schema_html}

<h2>6. Tool Catalog ({len(tools)} tools)</h2>
{tools_html}

<h2>7. Resource Catalog ({len(resources)} resources)</h2>
{resources_html}

<h2>8. How to Run</h2>
<pre>
# Install
cd sandbox/ltem-db-mcp
pip install -e .

# Configure
cp .env.example .env
# Edit .env with credentials

# Run
fastmcp run mcp_server/server.py:mcp --transport stdio
</pre>

<h2>9. Testing Checklist</h2>
<ul>
<li>[ ] <code>health_check</code> returns connected status</li>
<li>[ ] <code>list_tables</code> returns 3+ tables</li>
<li>[ ] <code>describe_table("ltem_historical_database")</code> returns columns</li>
<li>[ ] <code>get_regions</code> returns region list</li>
<li>[ ] <code>get_observations(region="Cabo Pulmo", limit=10)</code> returns rows</li>
<li>[ ] <code>calculate_diversity(region="La Paz")</code> returns indices</li>
<li>[ ] <code>compare_protection_levels</code> returns Kruskal-Wallis test</li>
<li>[ ] <code>trend_analysis(region="Cabo Pulmo")</code> returns trend statistics</li>
</ul>

<h2>10. Known Limitations</h2>
<ul>
<li>Environmental columns (SST, Chla) may not exist — tools handle gracefully</li>
<li>IDReef/IDSpecies type mismatches require CAST in joins</li>
<li>Biomass data may be sparse for some region/year combinations</li>
<li>No visualization output (by design — use Claude/chatMPA for rendering)</li>
<li>Schema cache is per-process (restart to refresh)</li>
</ul>

<h2>11. Framework Choice: FastMCP — Tradeoffs and Caveats</h2>

<h3>11.1 Why FastMCP Was Chosen</h3>
<ul>
<li><strong>Fast development:</strong> Minimal boilerplate, decorator-based tool registration, automatic JSON schema generation</li>
<li><strong>Protocol correctness:</strong> FastMCP handles MCP protocol details (stdio transport, JSON-RPC, capability negotiation)</li>
<li><strong>Reduced maintenance burden:</strong> Framework updates handle protocol changes, reducing need for manual protocol implementation</li>
<li><strong>Focus on domain logic:</strong> Allows concentration on LTEM querying and ecological analysis rather than MCP plumbing</li>
</ul>

<h3>11.2 Known Caveats and Limitations</h3>

<h4>Framework Lifecycle & Version Pinning</h4>
<ul>
<li>FastMCP is a relatively young framework with evolving APIs</li>
<li>Version pinned in <code>pyproject.toml</code> to ensure stability</li>
<li>Breaking changes in future FastMCP releases may require migration work</li>
<li>Dependency on upstream maintenance and security patches</li>
</ul>

<h4>Stdio Transport Limitations (Local-First, Not Multi-User Hosting)</h4>
<ul>
<li>FastMCP stdio transport is designed for single-user, local execution</li>
<li>Not suitable for centralized hosting or multi-tenant scenarios without additional infrastructure</li>
<li>Each client (Windsurf, chatMPA Studio) spawns its own MCP server process</li>
<li>No built-in request queuing, rate limiting, or connection pooling across users</li>
</ul>

<h4>Dependency Footprint as MCP Count Grows</h4>
<ul>
<li>Each MCP server is a separate Python package with its own dependencies</li>
<li>As chatMPA Studio adds more MCPs (e.g., spatial analysis, PDF parsing), total dependency count increases</li>
<li>Potential for version conflicts between MCPs sharing common libraries (pandas, numpy, etc.)</li>
<li>No shared runtime or dependency deduplication across MCPs</li>
</ul>

<h4>Need for Query Guardrails (Limits, Timeouts)</h4>
<ul>
<li>FastMCP does not enforce query complexity limits or timeouts by default</li>
<li>This server implements manual guardrails: 5000-row LIMIT cap, 20-second query timeout</li>
<li>Large result sets or expensive queries can still impact performance</li>
<li>No automatic query optimization or result pagination</li>
</ul>

<h4>Version/Config Management Across Many MCPs</h4>
<ul>
<li>Each MCP has its own <code>.env</code> file and configuration</li>
<li>No centralized config management for multiple MCPs</li>
<li>Credential rotation requires updating each MCP's <code>.env</code> independently</li>
<li>Version upgrades must be coordinated manually across MCPs</li>
</ul>

<h4>Client-Specific Integration Differences</h4>
<ul>
<li>MCP clients (Windsurf, Antigravity, Claude Desktop) may implement protocol features differently</li>
<li>Resource URIs, tool schemas, and error handling may behave inconsistently across clients</li>
<li>Testing required for each target client environment</li>
<li>No universal MCP testing framework or compatibility matrix</li>
</ul>

<h3>11.3 Mitigations Already in Place</h3>
<ul>
<li><strong>Version pinning:</strong> <code>pyproject.toml</code> locks FastMCP and all dependencies to known-good versions</li>
<li><strong>Read-only DB user:</strong> <code>mcp_ltem_ro</code> with SELECT-only privileges prevents accidental writes</li>
<li><strong>Parameterized queries:</strong> All user-facing tools use parameterized SQL to prevent injection</li>
<li><strong>Row limits and aggregation-first design:</strong> Tools return aggregated statistics rather than raw dumps where possible</li>
<li><strong>Local stdio execution model:</strong> No network exposure, credentials stay local, no shared state between users</li>
</ul>

<h3>11.4 Explicit Non-Goals (v1)</h3>
<ul>
<li><strong>No centralized hosting:</strong> Each user runs their own MCP server process locally</li>
<li><strong>No auto-scaling:</strong> Performance is bounded by single-process Python and database connection limits</li>
<li><strong>No cross-MCP orchestration:</strong> MCPs operate independently; no shared context or workflow coordination</li>
<li><strong>No write access:</strong> Read-only by design; data mutations happen outside the MCP layer</li>
<li><strong>No persistence or caching layer:</strong> Schema cache is in-memory and per-process only</li>
</ul>

<h3>11.5 Forward-Looking Note (Non-Committal)</h3>
<ul>
<li>FastMCP does not prevent future hosting or transport changes (e.g., HTTP/SSE, WebSocket)</li>
<li>Hosting and registry decisions are deferred intentionally to avoid premature architecture</li>
<li>MCPs are treated as versioned, installable packages that can be deployed in various configurations</li>
<li>Migration to alternative frameworks or custom protocol implementations remains possible if requirements change</li>
</ul>

</body>
</html>"""
	return html


def generate_docx(schema: dict | None, output_path: Path, version: str = "unknown", git_hash: str = "unknown") -> bool:
	"""Generate DOCX report. Returns True if successful."""
	try:
		from docx import Document
		from docx.shared import Inches, Pt
	except ImportError:
		print("  python-docx not installed. Skipping DOCX generation.")
		print("  Install with: pip install python-docx")
		return False

	doc = Document()
	timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

	doc.add_heading("LTEM Database MCP Server — Build Report", 0)
	doc.add_paragraph(f"Version: {version} | Git: {git_hash} | Generated: {timestamp}")

	doc.add_heading("1. Project Goal and Scope", level=1)
	doc.add_paragraph(
		"A FastMCP server connecting to the AWS RDS MySQL ecological_monitoring "
		"database, exposing the LTEM analysis suite as MCP tools."
	)

	doc.add_heading("2. Security Model", level=1)
	for item in [
		"Read-only database user (mcp_ltem_ro, SELECT only)",
		"Credentials via .env file (never hardcoded)",
		"SQL validation: only SELECT/SHOW/DESCRIBE allowed",
		"Table whitelist, auto LIMIT cap (5000), query timeout (20s)",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("3. Tool Catalog", level=1)
	tools = get_tool_catalog()
	table = doc.add_table(rows=1, cols=3)
	table.style = "Light List Accent 1"
	hdr = table.rows[0].cells
	hdr[0].text = "Module"
	hdr[1].text = "Tool"
	hdr[2].text = "Description"
	for t in tools:
		row = table.add_row().cells
		row[0].text = t["module"]
		row[1].text = t["name"]
		row[2].text = t["description"]

	doc.add_heading("4. Resource Catalog", level=1)
	resources = get_resource_catalog()
	table2 = doc.add_table(rows=1, cols=2)
	table2.style = "Light List Accent 1"
	hdr2 = table2.rows[0].cells
	hdr2[0].text = "URI"
	hdr2[1].text = "Description"
	for r in resources:
		row = table2.add_row().cells
		row[0].text = r["uri"]
		row[1].text = r["description"]

	doc.add_heading("5. Schema Discovery", level=1)
	if schema and "tables" in schema:
		for tbl, info in schema["tables"].items():
			doc.add_heading(f"{tbl} ({info.get('row_count', '?')} rows)", level=2)
			if "error" in info:
				doc.add_paragraph(f"Error: {info['error']}")
				continue
			cols = info.get("columns", [])
			if cols:
				tbl_doc = doc.add_table(rows=1, cols=3)
				tbl_doc.style = "Light List Accent 1"
				h = tbl_doc.rows[0].cells
				h[0].text = "Column"
				h[1].text = "Type"
				h[2].text = "Nullable"
				for col in cols:
					r = tbl_doc.add_row().cells
					r[0].text = col["name"]
					r[1].text = col["type"]
					r[2].text = "Yes" if col.get("nullable") else "No"
	else:
		doc.add_paragraph("Could not connect to database for schema discovery.")

	doc.add_heading("6. How to Run", level=1)
	doc.add_paragraph("pip install -e .")
	doc.add_paragraph("cp .env.example .env  # fill credentials")
	doc.add_paragraph("fastmcp run mcp_server/server.py:mcp --transport stdio")

	doc.add_heading("7. Framework Choice: FastMCP — Tradeoffs and Caveats", level=1)

	doc.add_heading("7.1 Why FastMCP Was Chosen", level=2)
	for item in [
		"Fast development: Minimal boilerplate, decorator-based tool registration, automatic JSON schema generation",
		"Protocol correctness: FastMCP handles MCP protocol details (stdio transport, JSON-RPC, capability negotiation)",
		"Reduced maintenance burden: Framework updates handle protocol changes, reducing need for manual protocol implementation",
		"Focus on domain logic: Allows concentration on LTEM querying and ecological analysis rather than MCP plumbing",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("7.2 Known Caveats and Limitations", level=2)

	doc.add_heading("Framework Lifecycle & Version Pinning", level=3)
	for item in [
		"FastMCP is a relatively young framework with evolving APIs",
		"Version pinned in pyproject.toml to ensure stability",
		"Breaking changes in future FastMCP releases may require migration work",
		"Dependency on upstream maintenance and security patches",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("Stdio Transport Limitations (Local-First, Not Multi-User Hosting)", level=3)
	for item in [
		"FastMCP stdio transport is designed for single-user, local execution",
		"Not suitable for centralized hosting or multi-tenant scenarios without additional infrastructure",
		"Each client (Windsurf, chatMPA Studio) spawns its own MCP server process",
		"No built-in request queuing, rate limiting, or connection pooling across users",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("Dependency Footprint as MCP Count Grows", level=3)
	for item in [
		"Each MCP server is a separate Python package with its own dependencies",
		"As chatMPA Studio adds more MCPs (e.g., spatial analysis, PDF parsing), total dependency count increases",
		"Potential for version conflicts between MCPs sharing common libraries (pandas, numpy, etc.)",
		"No shared runtime or dependency deduplication across MCPs",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("Need for Query Guardrails (Limits, Timeouts)", level=3)
	for item in [
		"FastMCP does not enforce query complexity limits or timeouts by default",
		"This server implements manual guardrails: 5000-row LIMIT cap, 20-second query timeout",
		"Large result sets or expensive queries can still impact performance",
		"No automatic query optimization or result pagination",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("Version/Config Management Across Many MCPs", level=3)
	for item in [
		"Each MCP has its own .env file and configuration",
		"No centralized config management for multiple MCPs",
		"Credential rotation requires updating each MCP's .env independently",
		"Version upgrades must be coordinated manually across MCPs",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("Client-Specific Integration Differences", level=3)
	for item in [
		"MCP clients (Windsurf, Antigravity, Claude Desktop) may implement protocol features differently",
		"Resource URIs, tool schemas, and error handling may behave inconsistently across clients",
		"Testing required for each target client environment",
		"No universal MCP testing framework or compatibility matrix",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("7.3 Mitigations Already in Place", level=2)
	for item in [
		"Version pinning: pyproject.toml locks FastMCP and all dependencies to known-good versions",
		"Read-only DB user: mcp_ltem_ro with SELECT-only privileges prevents accidental writes",
		"Parameterized queries: All user-facing tools use parameterized SQL to prevent injection",
		"Row limits and aggregation-first design: Tools return aggregated statistics rather than raw dumps where possible",
		"Local stdio execution model: No network exposure, credentials stay local, no shared state between users",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("7.4 Explicit Non-Goals (v1)", level=2)
	for item in [
		"No centralized hosting: Each user runs their own MCP server process locally",
		"No auto-scaling: Performance is bounded by single-process Python and database connection limits",
		"No cross-MCP orchestration: MCPs operate independently; no shared context or workflow coordination",
		"No write access: Read-only by design; data mutations happen outside the MCP layer",
		"No persistence or caching layer: Schema cache is in-memory and per-process only",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.add_heading("7.5 Forward-Looking Note (Non-Committal)", level=2)
	for item in [
		"FastMCP does not prevent future hosting or transport changes (e.g., HTTP/SSE, WebSocket)",
		"Hosting and registry decisions are deferred intentionally to avoid premature architecture",
		"MCPs are treated as versioned, installable packages that can be deployed in various configurations",
		"Migration to alternative frameworks or custom protocol implementations remains possible if requirements change",
	]:
		doc.add_paragraph(item, style="List Bullet")

	doc.save(str(output_path))
	return True


def main():
	docs_dir = PROJECT_ROOT / "docs"
	docs_dir.mkdir(exist_ok=True)

	version = get_version()
	git_hash = get_git_hash()

	print("LTEM Database MCP Server — Build Report Generator")
	print("=" * 50)
	print(f"  Version: {version}")
	print(f"  Git:     {git_hash}")

	# Try to get live schema
	print("\nAttempting database connection for schema discovery...")
	schema = get_schema_info()
	if schema:
		print("  Schema discovered successfully.")
	else:
		print("  Proceeding without live schema data.")

	# Generate HTML
	html_path = docs_dir / "build_report.html"
	print(f"\nGenerating HTML report: {html_path}")
	html_content = generate_html(schema, version=version, git_hash=git_hash)
	html_path.write_text(html_content, encoding="utf-8")
	print("  Done.")

	# Generate DOCX
	docx_path = docs_dir / "build_report.docx"
	print(f"\nGenerating DOCX report: {docx_path}")
	generate_docx(schema, docx_path, version=version, git_hash=git_hash)

	print("\nReport generation complete.")


if __name__ == "__main__":
	main()
